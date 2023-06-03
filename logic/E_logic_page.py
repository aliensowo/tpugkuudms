from PyQt5.QtWidgets import QTableWidgetItem, QFileDialog, QWidget

from gui.E_winner_page import Ui_MainWindow as EPage
from PyQt5 import QtWidgets
import docx
import os


class ELogicPage(QtWidgets.QMainWindow):

    def __init__(self, winner: dict, compare_id: int, username: str):
        super(ELogicPage, self).__init__()
        self.ui = EPage()
        self.ui.setupUi(self)
        self.winner = winner
        self.compare_id = compare_id
        self.username = username
        # button logic
        self.ui.pushButton.clicked.connect(self.button_logic)
        self.ui.pushButton_2.clicked.connect(self.close)

        # label logic
        self.ui.label_7.setText(self.winner["contractor_name"])
        self.viewTable()

    def viewTable(self):
        # 0;1
        item01 = QTableWidgetItem()
        item01.setText(self.winner["work_cost"].replace("\n", "").strip())
        self.ui.tableWidget.setItem(0, 1, item01)
        # 1;1
        item11 = QTableWidgetItem()
        item11.setText(self.winner["availability_of_technology"].replace("\n", "").strip())
        self.ui.tableWidget.setItem(1, 1, item11)
        # 2;1
        item21 = QTableWidgetItem()
        item21.setText(self.winner["availability_of_technology"].replace("\n", "").strip())
        self.ui.tableWidget.setItem(2, 1, item21)
        # 3;1
        item31 = QTableWidgetItem()
        item31.setText(self.winner["availability_of_technology"].replace("\n", "").strip())
        self.ui.tableWidget.setItem(3, 1, item31)
        # 4;1
        item41 = QTableWidgetItem()
        item41.setText(self.winner["availability_of_technology"].replace("\n", "").strip())
        self.ui.tableWidget.setItem(4, 1, item41)

    def button_logic(self):
        doc = docx.Document('raport1.docx')
        par_index = 0
        for paragraph in doc.paragraphs:
            style = paragraph.style
            if "{contractor_name}" in paragraph.text:
                paragraph.text = paragraph.text.format(
                    contractor_name=self.winner["contractor_name"].replace("\n", "").strip())
                doc.paragraphs[par_index].style = style
            elif "«_»" in paragraph.text:
                paragraph.text = paragraph.text = f"«{self.ui.dateEdit.date().day()}» {self.ui.dateEdit.date().month()} {self.ui.dateEdit.date().year()}г."
                doc.paragraphs[par_index].style = style
            par_index += 1
        doc.tables[0].columns[2].cells[1].text = self.winner["work_cost"].replace(
            "\n", "").strip()
        doc.tables[0].columns[2].cells[2].text = self.winner["availability_of_technology"].replace(
            "\n", "").strip()
        doc.tables[0].columns[2].cells[3].text = self.winner["tech_equipment"].replace(
            "\n", "").strip()
        doc.tables[0].columns[2].cells[4].text = self.winner["availability_of_production_facilities"].replace(
            "\n", "").strip()
        doc.tables[0].columns[2].cells[5].text = self.winner["qualified_human_resources_personnel"].replace(
            "\n", "").strip()
        result_file_name = self.saveFileDialog()
        if result_file_name:
            doc.save(result_file_name)
        else:
            doc.save(
                os.path.join(
                    f"report1_{self.username}_{self.compare_id}.docx"
                )
            )
            self.close()
        self.close()

    def saveFileDialog(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "QFileDialog.getSaveFileName()", "",
                                                  "All Files (*);;Text Files (*.docx)", options=options)
        if fileName:
            return fileName
        return None
