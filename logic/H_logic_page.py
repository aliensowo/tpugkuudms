import os
import time

from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QTableWidgetItem
import sqlalchemy.exc
from gui.H_work4_page import Ui_MainWindow as HPage
from models.database import SessionLocal
from logic._base_class_logic import BaseClassLogic
from models.crud import objects_contracts, neighborhood_directory, plan, fact, work
from typing import List
import docx


class HLogicPage(QtWidgets.QMainWindow, BaseClassLogic):
    docs = {
        1: {
            "name": "Справочник контрактов с датами начала и конца выполнения объекта по предмету договора",
            "fields": ["Номер объекта", "Наименование объекта", "Начало выполнения", "Конец выполнения"],
            "fields_m": ["id_object", "object_name", "date_start", "date_end"],
            "crud": objects_contracts,
            "get_curr": objects_contracts.list_current_y
        },
        2: {
            "name": "Данные контрактов с подрядными организациями по плановым срокам выполнения работ",
            "fields": ["Номер работы", "Наименование работ", "Плановый объем работ (м3)",
                       "Плановый срок начала проведения работ", "Плановый срок окончания проведения работ",
                       "Плановые сроки выполнения работ (дн.)"],
            "fields_m": ["id_work", "work_name", "volume",
                         "work_start",
                         "work_end",
                         "work_complete"],
            "crud": plan
        },
        3: {
            "name": "Данные по фактическим срокам выполнения работ на объекте и фактическим объемом работ",
            "fields": ["Номер работы", "Наименование работы", "Фактический объем работ (м3)",
                       "Фактический срок начала проведения работ", "Фактический срок окончания проведения работ",
                       "Фактические сроки выполнения работ (дн.)"],
            "fields_m": ["id_work", "work_name", "volume",
                         "work_start",
                         "work_end",
                         "work_complete"],
            "crud": fact
        }
    }

    objects_ids = []

    def __init__(self):
        super(HLogicPage, self).__init__()
        self.error_window = None
        self.winner_window = None
        self.ui = HPage()
        self.ui.setupUi(self)
        self.ui.pushButton.setVisible(False)
        self.ui.pushButton_2.setVisible(False)
        self.ui.tableWidget.setVisible(False)
        self.ui.pushButton_4.setVisible(False)
        self.ui.pushButton_3.clicked.connect(self.init_work_area)
        self.ui.pushButton_2.clicked.connect(self.prev_page)  # go back
        self.ui.pushButton_4.clicked.connect(self.save_data)  # save items
        self.ui.pushButton.clicked.connect(self.make_doc)  # save items
        self.ui.tableWidget.cellChanged.connect(self.cell_on_changed)
        self.ui.label_2.setVisible(False)
        self.ui.lineEdit.setVisible(False)

    def save_data(self):
        struct = self.get_doc_struct()
        column = self.ui.tableWidget.columnCount()
        row = self.ui.tableWidget.rowCount()
        with SessionLocal() as session:
            for i in range(row):
                if self.check_row_4_empty(i, column):
                    data = []
                    for j in range(column):
                        item = self.ui.tableWidget.item(i, j)
                        data.append(item.text())
                    create = struct["crud"].__getattribute__("create")
                    try:
                        create(session, *data)
                    except sqlalchemy.exc.IntegrityError:
                        continue
                    except sqlalchemy.exc.PendingRollbackError:
                        continue
                    if struct["crud"] in [plan, fact]:
                        try:
                            work.create(session, *data[:2])
                        except sqlalchemy.exc.IntegrityError:
                            continue

    def check_row_4_empty(self, row, column):
        for i in range(column):
            item = self.ui.tableWidget.item(row, i)
            if item is None:
                return False
        return True

    def init_work_area(self):
        self.build_table()
        if self.ui.pushButton_3.text() == "Начать":
            self.ui.pushButton_3.setText("Далее")
            self.ui.pushButton_2.setVisible(True)
            self.ui.pushButton_4.setVisible(True)

    def build_table(self):
        cur_key = self.get_text_label()
        if cur_key == "":
            self.ui.tableWidget.setVisible(True)
            self.ui.label.setText(self.docs[1]["name"])
            self.__build_table()
            self.ui.label_2.setVisible(True)
            self.ui.lineEdit.setVisible(True)
        else:
            object_ids = self.ui.lineEdit.text().split(",")
            if object_ids:
                self.ui.lineEdit.setText("")
                self.objects_ids = object_ids
            self.ui.label_2.setVisible(False)
            self.ui.lineEdit.setVisible(False)
            np = self.get_next_page()
            if np:
                self.ui.label.setText(self.get_next_page())
                self.clear_table()
                self.__build_table()
            else:
                self.ui.pushButton.setVisible(True)

    def __build_table(self):
        current = self.get_doc_struct()
        with SessionLocal() as session:
            try:
                if current["name"] == "Справочник контрактов с датами начала и конца выполнения объекта по предмету договора":
                    listq = current["get_curr"]
                    query_array = listq(session)
                else:
                    listq = current["crud"].__getattribute__("list")
                    query_array = listq(session)
            except KeyError:
                query_array = []
        if query_array.__len__() == 0:
            self.ui.tableWidget.setRowCount(1)
        else:
            self.ui.tableWidget.setRowCount(query_array.__len__() + 1)
        self.ui.tableWidget.setColumnCount(current["fields"].__len__())
        self.ui.tableWidget.setHorizontalHeaderLabels(current["fields"])
        header = self.ui.tableWidget.horizontalHeader()
        header.setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)
        if query_array:
            for row in range(len(query_array)):
                row_elem: dict = query_array[row]
                for column in range(len(current["fields_m"])):
                    item = QTableWidgetItem()
                    item.setText(str(row_elem[current["fields_m"][column]]))
                    self.ui.tableWidget.setItem(row, column, item)

    def prev_page(self):
        pp = self.get_prev_page()
        self.ui.tableWidget.clear()
        if pp:
            self.ui.label.setText(pp)
            self.build_table()
        else:
            self.ui.label.setText("")
            self.ui.pushButton_2.setVisible(False)
            self.ui.tableWidget.setVisible(False)
            self.ui.pushButton_3.setText("Начать")

    def clear_table(self):
        self.ui.tableWidget.clear()

    def get_text_label(self):
        return self.ui.label.text()

    def get_doc_struct(self):
        key = self.get_text_label()
        for k, v in self.docs.items():
            if v["name"] == key:
                return self.docs[k]

    def get_next_page(self):
        cur_key = self.get_text_label()
        for k, v in self.docs.items():
            if v["name"] == cur_key:
                try:
                    return self.docs[k + 1]["name"]
                except KeyError:
                    return 0

    def get_prev_page(self):
        cur_key = self.get_text_label()
        for k, v in self.docs.items():
            if v["name"] == cur_key:
                try:
                    return self.docs[k - 2]["name"]
                except KeyError:
                    return 0

    def cell_on_changed(self):
        cell = self.ui.tableWidget.currentItem()
        if cell is not None:
            row = cell.row()
            table_column = self.ui.tableWidget.columnCount()
            if self.check_row_4_empty(row, table_column):
                self.ui.tableWidget.insertRow(row + 1)

    def make_doc(self):
        itog = []
        with SessionLocal() as session:
            m = objects_contracts.list_current_y(session)
            for mo in m:
                itog_row = {"Код объекта, подлежащего сдаче": mo["id_object"],
                            "Наименование объекта": mo["object_name"]}
                cv_plan = []
                cv_fact = []

                plan_list: List[plan.models.Plan] = plan.list(session)
                for p in plan_list:
                    if self.check_range(p["work_start"], p["work_end"], mo["date_start"], mo["date_end"]):
                        cv_plan.append(p)
                fact_list: List[fact.models.Fact] = fact.list(session)
                for f in fact_list:
                    if self.check_range(f["work_start"], f["work_end"], mo["date_start"], mo["date_end"]):
                        cv_fact.append(f)
                try:
                    plan_sroc_percent = 100 / sum([int(x["work_complete"]) for x in cv_plan]) * sum(
                        [int(x["work_complete"]) for x in cv_fact])
                except ZeroDivisionError:
                    plan_sroc_percent = 0
                try:
                    plan_volume_percent = 100 / sum([int(x["volume"]) for x in cv_plan]) * sum(
                        [int(x["volume"]) for x in cv_fact])
                except ZeroDivisionError:
                    plan_volume_percent = 0
                if plan_sroc_percent in range(95, 105) and plan_volume_percent in range(97, 103):
                    itog_row["Заключение"] = "Подлежит для сдачи в текущем году"
                else:
                    itog_row["Заключение"] = "Не подлежит для сдачи в текущем году"
                itog_row["Дата сдачи объекта"] = mo["date_end"]
                itog.append(itog_row)
        doc = docx.Document('raport4.docx')
        for row_itog in itog:
            index = itog.index(row_itog) + 1
            doc.tables[0].columns[0].cells[index].text = str(row_itog["Код объекта, подлежащего сдаче"])
            doc.tables[0].columns[1].cells[index].text = str(row_itog["Наименование объекта"])
            doc.tables[0].columns[2].cells[index].text = str(row_itog["Заключение"])
            doc.tables[0].columns[3].cells[index].text = str(row_itog["Дата сдачи объекта"])
        try:
            doc.save(
                os.path.join(
                    os.path.join(os.getenv("USERPROFILE")),
                    "Desktop",
                    f"report4_itog.docx"
                )
            )
        except Exception as e:
            doc.save(
                os.path.join(
                    # os.path.join(os.getenv("USERPROFILE")),
                    # "Desktop",
                    f"report4_itog.docx"
                )
            )
        self.close()

    @staticmethod
    def str_date_to_sec(string_date: str):
        return int(time.mktime(time.strptime(string_date.strip(), "%d.%m.%Y")))

    @staticmethod
    def check_range(es, ee, ts, te):
        if HLogicPage.str_date_to_sec(es) in range(
                HLogicPage.str_date_to_sec(ts), HLogicPage.str_date_to_sec(te)
        ) and HLogicPage.str_date_to_sec(ee) in range(
            HLogicPage.str_date_to_sec(ts), HLogicPage.str_date_to_sec(te)
        ):
            return True
