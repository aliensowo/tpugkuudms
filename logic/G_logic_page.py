import os
import time

import docx
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtGui import QColor
from PyQt5.QtWidgets import QTableWidgetItem
import sqlalchemy.exc
from gui.G_work3_page import Ui_MainWindow as GPage
from models.database import SessionLocal
from models.crud import contract_values, act_ks_6a, act_ks_2
from logic._base_class_logic import BaseClassLogic

from typing import List
import math


class GLogicPage(QtWidgets.QMainWindow, BaseClassLogic):
    docs = {
        1: {
            "name": "Данные контрактов с подрядными организациями по плановым объемам работ",
            "fields": ["Код по видам работ", "Наименование вида работ", "Плановый объем работ (м3)"],
            "fields_m": ["id_object", "object_name", "work_name", "work_value"],
            "crud": contract_values,
        },
        2: {
            "name": "Акт выполнения работ (КС-6а)",
            "fields": ["Код по видам работ", "Номер по порядку", "Номер позиции по смете",
                       "Номер единичной расценки", "Единица измерения",
                       "Цена за единицу, руб.", "Количество работ по смете", "Сметная (договорная) стоимость, руб.",
                       "Месяц выполнения работ", "количество", "стоимость, руб.",
                       "стоимость фактически выполненных работ с начала строительства, руб."],
            "fields_m": ["id_work", "number", "smeta_number"
                                              "unit_number", "measure",
                         "cost_for_unit", "work_count_smeta", "smeta_cost",
                         "month", "count", "cost",
                         "cost_fact"],
            "crud": act_ks_6a
        },
        3: {
            "name": "Акт выполнения работ (КС-2)",
            "fields": ["Код по видам работ", "Номер по порядку", "Номер позиции по смете",
                       "Номер единичной расценки", "Единица измерения",
                       "Выполнено работ количество", "Выполнено работ цена за единицу", "Выполнено работ стоимость"],
            "fields_m": ["id_work", "number", "smeta_number"
                                              "unit_number", "measure",
                         "work_count", "work_unit", "work_cost"],
            "crud": act_ks_2
        }
    }

    def __init__(self):
        super(GLogicPage, self).__init__()
        self.error_window = None
        self.winner_window = None
        self.ui = GPage()
        self.ui.setupUi(self)
        self.ui.pushButton.setVisible(False)
        self.ui.pushButton_2.setVisible(False)
        self.ui.tableWidget.setVisible(False)
        self.ui.pushButton_4.setVisible(False)
        self.ui.pushButton_3.clicked.connect(self.init_work_area)
        self.ui.pushButton_2.clicked.connect(self.prev_page)  # go back
        self.ui.pushButton_4.clicked.connect(self.save_data)  # save items
        self.ui.pushButton.clicked.connect(self.make_doc)  # save items
        self.ui.tableWidget.cellChanged.connect(self.cell_on_changed)
        self.ui.label_2.setVisible(False)
        self.ui.lineEdit.setVisible(False)

    def save_data(self):
        struct = self.get_doc_struct()
        column = self.ui.tableWidget.columnCount()
        row = self.ui.tableWidget.rowCount()
        with SessionLocal() as session:
            for i in range(row):
                if self.check_row_4_empty(i, column):
                    data = []
                    for j in range(column):
                        item = self.ui.tableWidget.item(i, j)
                        data.append(item.text())
                    create = struct["crud"].__getattribute__("create")
                    try:
                        create(session, *data)
                    except sqlalchemy.exc.IntegrityError:
                        continue

    def check_row_4_empty(self, row, column):
        for i in range(column):
            item = self.ui.tableWidget.item(row, i)
            if item is None:
                return False
        return True

    def init_work_area(self):
        self.build_table()
        if self.ui.pushButton_3.text() == "Начать":
            self.ui.pushButton_3.setText("Далее")
            self.ui.pushButton_2.setVisible(True)
            self.ui.pushButton_4.setVisible(True)

    def build_table(self):
        cur_key = self.get_text_label()
        if cur_key == "":
            self.ui.tableWidget.setVisible(True)
            self.ui.label.setText(self.docs[1]["name"])
            self.__build_table()
            self.ui.label_2.setVisible(True)
            self.ui.lineEdit.setVisible(True)
        else:
            object_ids = self.ui.lineEdit.text().split(",")
            if object_ids:
                self.ui.lineEdit.setText("")
                self.objects_ids = object_ids
            self.ui.label_2.setVisible(False)
            self.ui.lineEdit.setVisible(False)
            np = self.get_next_page()
            if np:
                self.ui.label.setText(self.get_next_page())
                self.clear_table()
                self.__build_table()
            else:
                self.ui.pushButton.setVisible(True)

    def __build_table(self):
        current = self.get_doc_struct()
        with SessionLocal() as session:
            try:
                listq = current["crud"].__getattribute__("list")
                query_array = listq(session)
            except KeyError:
                query_array = []
        if query_array.__len__() == 0:
            self.ui.tableWidget.setRowCount(1)
        else:
            self.ui.tableWidget.setRowCount(query_array.__len__() + 1)
        self.ui.tableWidget.setColumnCount(current["fields"].__len__())
        self.ui.tableWidget.setHorizontalHeaderLabels(current["fields"])
        header = self.ui.tableWidget.horizontalHeader()
        header.setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)
        header.setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)
        if query_array:
            for row in range(len(query_array)):
                row_elem: dict = query_array[row]
                for column in range(len(current["fields_m"])):
                    item = QTableWidgetItem()
                    item.setText(str(row_elem[current["fields_m"][column]]))
                    self.ui.tableWidget.setItem(row, column, item)

    def prev_page(self):
        pp = self.get_prev_page()
        self.ui.tableWidget.clear()
        if pp:
            self.ui.label.setText(pp)
            self.build_table()
        else:
            self.ui.label.setText("")
            self.ui.pushButton_2.setVisible(False)
            self.ui.tableWidget.setVisible(False)
            self.ui.pushButton_3.setText("Начать")

    def clear_table(self):
        self.ui.tableWidget.clear()

    def get_text_label(self):
        return self.ui.label.text()

    def get_doc_struct(self):
        key = self.get_text_label()
        for k, v in self.docs.items():
            if v["name"] == key:
                return self.docs[k]

    def get_next_page(self):
        cur_key = self.get_text_label()
        for k, v in self.docs.items():
            if v["name"] == cur_key:
                try:
                    return self.docs[k + 1]["name"]
                except KeyError:
                    return 0

    def get_prev_page(self):
        cur_key = self.get_text_label()
        for k, v in self.docs.items():
            if v["name"] == cur_key:
                try:
                    return self.docs[k - 2]["name"]
                except KeyError:
                    return 0

    def cell_on_changed(self):
        cell = self.ui.tableWidget.currentItem()
        if cell is not None:
            row = cell.row()
            table_column = self.ui.tableWidget.columnCount()
            if self.check_row_4_empty(row, table_column):
                self.ui.tableWidget.insertRow(row + 1)

    def make_doc(self):
        itog = []
        with SessionLocal() as session:
            cv = contract_values.list(session)
            summary = 0
            for c in cv:
                itog_row = {}
                plan_value_of_work = c["work_value"]
                code_work = c["id_object"]
                work_name = c["object_name"]
                cs2 = act_ks_2.get_by_code_work(session, code_work)
                fact_value_of_work = cs2.work_count
                act_6a_list = act_ks_6a.get_by_code_work(session, code_work)
                grade_result = 0
                grade_cost = 0
                for act_6a in act_6a_list:
                    grade_result += int(act_6a["count"])
                    grade_cost += int(act_6a["cost_fact"])
                ostatok = abs(int(fact_value_of_work) - int(plan_value_of_work))
                summary += grade_cost

                itog_row["Id по видам работ"] = code_work
                itog_row["Вид выполненных работ"] = work_name
                itog_row["Единица измерения"] = c["measure"]
                itog_row["Всего по договору (плановый)"] = plan_value_of_work
                itog_row["За отчетный период (фактический)"] = fact_value_of_work
                itog_row["С нарастающим итогом объем"] = grade_result
                itog_row["Остаток"] = ostatok
                itog_row["С нарастающим итогом стоимость"] = grade_cost

                itog.append(itog_row)
        doc = docx.Document('raport3.docx')
        for row_itog in itog:
            index = itog.index(row_itog) + 1
            doc.tables[0].columns[0].cells[index + 1].text = str(itog_row["Id по видам работ"])
            doc.tables[0].columns[1].cells[index + 1].text = str(itog_row["Вид выполненных работ"])
            doc.tables[0].columns[2].cells[index + 1].text = str(itog_row["Единица измерения"])
            doc.tables[0].columns[3].cells[index + 1].text = str(itog_row["Всего по договору (плановый)"])
            doc.tables[0].columns[4].cells[index + 1].text = str(itog_row["За отчетный период (фактический)"])
            doc.tables[0].columns[5].cells[index + 1].text = str(itog_row["С нарастающим итогом объем"])
            doc.tables[0].columns[6].cells[index + 1].text = str(itog_row["Остаток"])
            doc.tables[0].columns[7].cells[index + 1].text = str(itog_row["С нарастающим итогом стоимость"])
        doc.tables[0].columns[8].cells[2].text = str(summary)
        try:
            doc.save(
                os.path.join(
                    os.path.join(os.getenv("USERPROFILE")),
                    "Desktop",
                    f"report3_itog.docx"
                )
            )
        except Exception as e:
            doc.save(
                os.path.join(
                    # os.path.join(os.getenv("USERPROFILE")),
                    # "Desktop",
                    f"report3_itog.docx"
                )
            )
        self.close()

    @staticmethod
    def str_date_to_sec(string_date: str):
        return int(time.mktime(time.strptime(string_date.strip(), "%d.%m.%Y")))

    @staticmethod
    def check_range(es, ee, ts, te):
        if GLogicPage.str_date_to_sec(es) in range(
                GLogicPage.str_date_to_sec(ts), GLogicPage.str_date_to_sec(te)
        ) and GLogicPage.str_date_to_sec(ee) in range(
            GLogicPage.str_date_to_sec(ts), GLogicPage.str_date_to_sec(te)
        ):
            return True
